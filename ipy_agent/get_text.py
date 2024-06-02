import os
_root_=os.path.dirname(os.path.abspath(__file__))
import sys
if not sys.path[0]==_root_:
    sys.path.insert(0,_root_)
def root_join(*args):
    return os.path.join(_root_,*args)

from bs4 import BeautifulSoup
import PyPDF2
import docx
import odf
import os
import requests
from io import BytesIO
import json
import inspect
from get_webdriver import get_webdriver

def strip_newlines(string):
    while len(string)>=1 and string[0]=='\n':
        string=string[1:]
    
    while len(string)>=1 and string[-1]=='\n':
        string=string[:len(string)-1]
    
    while not (newstring:=string.replace('\n\n', '\n'))==string:
        string=newstring
    return string

def handle_directory(path):
    import os

    def recurse_folder(folder, prefix=''):
        contents = os.listdir(folder)
        output = ''
        for i, item in enumerate(contents):
            output += prefix + '├── ' + item + '\n'
            path = os.path.join(folder, item)
            if os.path.isdir(path):
                if i == len(contents) - 1:
                    output += recurse_folder(path, prefix + '    ')
                else:
                    output += recurse_folder(path, prefix + '│   ')
        return output

    if os.path.exists(path) and os.path.isdir(path):
        return path + '\n' + recurse_folder(path)
    else:
        return f'The path {path} is not a valid directory.'

def extract_webpage_content(url):
    driver = get_webdriver()
    driver.get(url)
    page_source = driver.page_source
    soup = BeautifulSoup(page_source, 'html.parser')
    text = soup.get_text()
    driver.quit()
    return strip_newlines(text)
 
def handle_class(source):
    class_info = {
        "type": "class",
        "name": source.__name__,
        "doc": inspect.getdoc(source) or "No documentation available.",
        "base_classes": [base.__name__ for base in inspect.getmro(source)[1:]],
        "methods": {},
        "class_methods": {},
        "static_methods": {},
        "properties": {},
        "attributes": {},
        "string_repr": repr(source)
    }
    # Process methods, class methods, static methods, and properties
    for name, member in inspect.getmembers(source):
        if inspect.isfunction(member):
            class_info["methods"][name] = {
                "doc": inspect.getdoc(member) or "No documentation available.",
                "signature": str(inspect.signature(member))
            }
        elif isinstance(member, classmethod):
            class_info["class_methods"][name] = {
                "doc": inspect.getdoc(member) or "No documentation available.",
                "signature": str(inspect.signature(member.__func__))
            }
        elif isinstance(member, staticmethod):
            class_info["static_methods"][name] = {
                "doc": inspect.getdoc(member) or "No documentation available.",
                # Static methods do not have a signature
            }
        elif isinstance(member, property):
            class_info["properties"][name] = {
                "doc": inspect.getdoc(member) or "No documentation available.",
                # Properties do not have a signature
            }
    # Process attributes
    for name in dir(source):
        if not name.startswith('__') and not inspect.isroutine(getattr(source, name)):
            attribute = inspect.getattr_static(source, name)
            try:
                attribute=json.dumps(attribute)
            except:
                attribute=repr(attribute)
            class_info["attributes"][name] = {
                "value": attribute,
                "type": type(attribute).__name__
            }
    return json.dumps(class_info, indent=4)

def handle_function(source):
    # Determine the specific type of callable
    if inspect.isfunction(source):
        callable_type = "function"
    elif inspect.ismethod(source):
        callable_type = "instance method"
    elif inspect.isbuiltin(source):
        callable_type = "built-in"
    elif inspect.ismethoddescriptor(source):
        callable_type = "method descriptor"
    else:
        callable_type = "other callable"

    callable_info = {
        "type": callable_type,
        "name": getattr(source, '__name__', 'Unnamed'),
        "doc": inspect.getdoc(source) or "No documentation available.",
        "module": getattr(source, '__module__', None),
        "string_repr": repr(source),
        "signature": None,
        "source": None
    }

    # Attempt to get the signature
    try:
        signature = inspect.signature(source)
        callable_info["signature"] = str(signature)
    except (TypeError, ValueError):
        callable_info["signature"] = "Not available"

    # Attempt to get the source code if it's a regular function
    if callable_type == "function":
        try:
            callable_info["source"] = inspect.getsource(source)
        except Exception:
            callable_info["source"] = "Not available"

    return json.dumps(callable_info, indent=4)

def handle_module(source):
    module_info = {
        "type": "module",
        "name": source.__name__,
        "doc": inspect.getdoc(source) or "No documentation available.",
        "functions": {},
        "classes": {},
        "variables": {},
        "submodules": {},
        "string_repr": repr(source)
    }
    for name, member in inspect.getmembers(source):
        if inspect.isfunction(member) and not name.startswith('_'):
            module_info["functions"][name] = {
                "doc": inspect.getdoc(member) or "No documentation available.",
                "signature": str(inspect.signature(member))
            }
        elif inspect.isclass(member) and not name.startswith('_'):
            module_info["classes"][name] = {
                "doc": inspect.getdoc(member) or "No documentation available."
            }
        elif inspect.ismodule(member) and not name.startswith('_'):
            module_info["submodules"][name] = {
                "doc": inspect.getdoc(member) or "No documentation available."
            }
        elif not (inspect.isroutine(member) or inspect.isclass(member) or inspect.ismodule(member)) and not name.startswith('_'):
            try:
                value=json.dumps(member)
            except:
                value=repr(member)
            module_info["variables"][name] = {
                "value": value,
                "type": type(member).__name__
            }
    return json.dumps(module_info, indent=4)

def handle_object(source):
    serialized=repr(source)
    if hasattr(source, 'dumps') and callable(getattr(source, 'dumps')):
        try:
            serialized=source.dumps()
        except:
            pass
    else:
        try:
            serialized=json.dumps(source)
        except:
            pass

    instance_info = {
        "type": "class_instance",
        "class": source.__class__.__name__,
        "doc": inspect.getdoc(source) or "No documentation available.",
        "attributes": {},
        "methods": {},
        "properties": {},
        "callable": callable(source),
        "call_signature": None,
        "string_repr": serialized
    }

    if instance_info["callable"]:
        try:
            instance_info["call_signature"] = str(inspect.signature(source.__call__))
        except (TypeError, ValueError):
            instance_info["call_signature"] = "Not available"

    # Use inspect.getmembers to safely retrieve object members
    members = inspect.getmembers(source, lambda a: not(inspect.isroutine(a)))
    methods = inspect.getmembers(source, inspect.isroutine)
    properties = [m for m in members if isinstance(m[1], property)]

    # Filter attributes, methods, and properties
    for name, member in members:
        if not name.startswith('__'):
            try:
                value=json.dumps(member)
            except:
                value=repr(member)
            instance_info["attributes"][name] = {
                "value": value,
                "type": type(member).__name__,
                "string_repr": repr(member)
            }

    for name, method in methods:
        if not name.startswith('__'):
            try:
                signature = str(inspect.signature(method))
            except (TypeError, ValueError):
                signature = "Not available"
            instance_info["methods"][name] = {
                "doc": inspect.getdoc(method) or "No documentation available.",
                "signature": signature
            }

    for name, prop in properties:
        instance_info["properties"][name] = {
            "doc": inspect.getdoc(prop) or "No documentation available."
        }

    # Serialize the instance information
    return json.dumps(instance_info, indent=4)    

def handle_url(source):
    url = source
    try:
        response = requests.get(url)
        content_type = response.headers['content-type']
        if content_type == 'application/pdf':
            with BytesIO(response.content) as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num in range(len(reader.pages)):
                    text += reader.pages[page_num].extract_text()
        elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            with BytesIO(response.content) as file:
                doc = docx.Document(file)
                text = "\n".join([para.text for para in doc.paragraphs])
        elif content_type == 'application/vnd.oasis.opendocument.text':
            with BytesIO(response.content) as file:
                doc = odf.opendocument.load(file)
                allparas = doc.getElementsByType(odf.text.P)
                text='\n'.join([odf.teletype.extractText(para) for para in allparas])
        else:
            text = extract_webpage_content(url)
        return text
    except requests.exceptions.RequestException as e:
        return f"Unable to process url. Connexion error : {e}"

def handle_file(source):
    ext = os.path.splitext(source)[1]
    try:
        if ext == '.pdf':
            with open(source, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num in range(len(reader.pages)):
                    text += reader.pages[page_num].extract_text()
        elif ext == '.docx':
            doc = docx.Document(source)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif ext == '.odt':
            doc = odf.opendocument.load(source)
            allparas = doc.getElementsByType(odf.text.P)
            text='\n'.join([odf.teletype.extractText(para) for para in allparas])
        elif ext == '.html':
            with open(source, 'r') as f:
                soup = BeautifulSoup(f, 'html.parser')
                text = soup.get_text()
        else:
            with open(source, 'r', encoding='utf-8') as f:
                text = f.read()
        return text
    except FileNotFoundError:
        return f"Unable to process file. File not found : {source}"
    except Exception as e:
        return f"Error while attempting to read file : {e}"

def get_text(source):
    if not isinstance(source,str):
        # Then check for classes
        if inspect.isclass(source):
            text=handle_class(source)
        # Check for callables (functions, methods, builtins)
        elif inspect.isfunction(source) or inspect.ismethod(source) or inspect.isbuiltin(source) or inspect.ismethoddescriptor(source):
            text=handle_function(source)                
        # Check for modules
        elif inspect.ismodule(source):
            text=handle_module(source)
        # Check for built-in types
        elif isinstance(source,(list,dict,tuple,set,int,float,bool)):
            try:
                text=json.dumps(source,indent=4)
            except:
                text=repr(source)
        # Finally, check for generic objects
        elif isinstance(source, object):
            text=handle_object(source)
        # Fallback to repr otherwise
        else:
            text=repr(source)

    elif source.startswith('http'):
        text=handle_url(source)

    elif os.path.isfile(source):
        text=handle_file(source)

    elif os.path.isdir(source):
        text=handle_directory(source)

    else:
        text = source
        
    return strip_newlines(text)